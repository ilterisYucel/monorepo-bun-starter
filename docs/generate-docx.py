#!/usr/bin/env python3
"""Markdown → DOCX converter with cover page, header, and embedded images."""

import re, os, argparse
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml.ns import qn
from docx.oxml import parse_xml, OxmlElement

MARKDOWN_PATH = os.path.join(os.path.dirname(__file__), "proje-tanitim-raporu.md")
LOGO_PATH = os.path.join(os.path.dirname(__file__), "logo.png")
ARCH_IMG_PATH = os.path.join(os.path.dirname(__file__), "images", "architecture.png")
OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "GD-ESS-EYS-Yazilim-Tanitim-Dokumani.docx")

COVER_TITLE = "GD-ESS EYS\nYazilim Tanitim Dokumani"
COVER_SUBTITLE = "Enerji Depolama Sistemi (EDS) Izleme, Kontrol ve Yonetim Yazilimi"
COVER_DATE = "Temmuz 2026  ·  Versiyon 1.0"

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(12)
HEADER_LOGO_WIDTH = Inches(6.27)   # full page width (A4: 21cm - 2.5cm*2 margins ≈ 16cm ≈ 6.3")
COVER_LOGO_WIDTH = Inches(3.5)


WML_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def set_cell_shading(cell, color_hex):
    """Set table cell background color."""
    tcPr = cell._tc.get_or_add_tcPr()
    shading = parse_xml(
        f'<w:shd xmlns:w="{WML_NS}" w:fill="{color_hex}" w:val="clear"/>'
    )
    tcPr.append(shading)


def set_table_borders(table):
    """Set borders for an entire table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr xmlns:w="{WML_NS}"/>')
    borders = parse_xml(
        f'<w:tblBorders xmlns:w="{WML_NS}">'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="333333"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="333333"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="333333"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="333333"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="333333"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="333333"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def add_header_image(section):
    """Add full-width logo to section header."""
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(LOGO_PATH, width=HEADER_LOGO_WIDTH)


def add_cover_page(doc):
    """Create a cover page with logo and title."""
    # Add empty paragraphs to push content down
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Logo centered
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(LOGO_PATH, width=COVER_LOGO_WIDTH)

    # Spacing
    for _ in range(2):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(COVER_TITLE)
    run.font.name = FONT_NAME
    run.font.size = Pt(26)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    # Spacing
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Subtitle line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(COVER_SUBTITLE)
    run.font.name = FONT_NAME
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    run = p.add_run(COVER_DATE)
    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(130, 130, 130)


def extract_headings():
    """Pre-scan markdown for headings (after metadata section) to build TOC."""
    with open(MARKDOWN_PATH, 'r', encoding='utf-8') as f:
        content = f.read()
    parts = content.split('\n---\n', 1)
    body = parts[1] if len(parts) > 1 else parts[0]

    headings = []
    for line in body.split('\n'):
        stripped = line.strip()
        if stripped.startswith('# ') and not stripped.startswith('## '):
            headings.append((1, stripped[2:].strip()))
        elif stripped.startswith('## ') and not stripped.startswith('### '):
            headings.append((2, stripped[3:].strip()))
        elif stripped.startswith('### '):
            headings.append((3, stripped[4:].strip()))
    return headings


def add_toc_page(doc, headings):
    """Build a real Word TOC field page. Cached entries are visible immediately;
    Word refreshes page numbers on open (dirty field) or on F9."""
    # Top spacing
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ICINDEKILER")
    run.font.name = FONT_NAME
    run.font.size = Pt(20)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    # Spacing
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)

    # TOC field begin
    toc_p = doc.add_paragraph()
    r = toc_p.add_run()
    fld = OxmlElement('w:fldChar')
    fld.set(qn('w:fldCharType'), 'begin')
    fld.set(qn('w:dirty'), 'true')
    r._element.append(fld)
    r = toc_p.add_run()
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    r._element.append(instr)
    r = toc_p.add_run()
    fld = OxmlElement('w:fldChar')
    fld.set(qn('w:fldCharType'), 'separate')
    r._element.append(fld)

    # Cached entries (one paragraph per entry, Word style)
    idx = 0
    for level, text in headings:
        if level == 1:
            continue  # skip H1 (document title, already on cover)

        anchor = f"_Toc{idx}"
        idx += 1

        entry_p = doc.add_paragraph()
        entry_p.paragraph_format.space_after = Pt(2)
        entry_p.paragraph_format.line_spacing = 1.4

        if level == 2:
            entry_p.paragraph_format.space_before = Pt(8)
            half_pts = 24
            bold = True
        else:
            entry_p.paragraph_format.left_indent = Cm(1.2)
            entry_p.paragraph_format.space_before = Pt(2)
            half_pts = 22
            bold = False

        hl = OxmlElement('w:hyperlink')
        hl.set(qn('w:anchor'), anchor)
        hl.set(qn('w:history'), '1')

        # Visible entry text run
        run_elem = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rStyle = OxmlElement('w:rStyle')
        rStyle.set(qn('w:val'), 'Hyperlink')
        rPr.append(rStyle)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(half_pts))
        rPr.append(sz)
        if bold:
            rPr.append(OxmlElement('w:b'))
        run_elem.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        run_elem.append(t)
        hl.append(run_elem)

        # Hidden tab run
        tab_r = OxmlElement('w:r')
        tab_rPr = OxmlElement('w:rPr')
        tab_rPr.append(OxmlElement('w:webHidden'))
        tab_r.append(tab_rPr)
        tab_r.append(OxmlElement('w:tab'))
        hl.append(tab_r)

        # PAGEREF field (page number), cached value
        for kind, payload in (
            ("begin", None),
            ("instr", f" PAGEREF {anchor} \\h "),
            ("separate", None),
            ("text", "3"),
            ("end", None),
        ):
            fr = OxmlElement('w:r')
            frPr = OxmlElement('w:rPr')
            frPr.append(OxmlElement('w:webHidden'))
            fr.append(frPr)
            if kind == "instr":
                it = OxmlElement('w:instrText')
                it.set(qn('xml:space'), 'preserve')
                it.text = payload
                fr.append(it)
            elif kind == "text":
                wt = OxmlElement('w:t')
                wt.text = payload
                fr.append(wt)
            else:
                fc = OxmlElement('w:fldChar')
                fc.set(qn('w:fldCharType'), kind)
                fr.append(fc)
            hl.append(fr)

        entry_p._element.append(hl)

    # TOC field end
    end_p = doc.add_paragraph()
    r = end_p.add_run()
    fld = OxmlElement('w:fldChar')
    fld.set(qn('w:fldCharType'), 'end')
    r._element.append(fld)

    # Page break after TOC
    doc.add_page_break()


def add_heading_bookmark(paragraph, bookmark_name, bookmark_id):
    """Insert a Word bookmark into a heading paragraph so TOC hyperlinks can target it."""
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(bookmark_id))
    start.set(qn('w:name'), bookmark_name)

    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(bookmark_id))

    paragraph._element.insert(0, start)
    paragraph._element.append(end)


LINK_RE = re.compile(r'(\[[^\]]*\]\([^)\s]+\)|https?://[^\s]+)')


def add_hyperlink_run(paragraph, url, text, font_size=None, bold=False):
    """Add a clickable external hyperlink run to a paragraph."""
    r_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hl = OxmlElement('w:hyperlink')
    hl.set(qn('r:id'), r_id)
    hl.set(qn('w:history'), '1')
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    if font_size is not None:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(font_size.pt * 2)))
        rPr.append(sz)
    if bold:
        rPr.append(OxmlElement('w:b'))
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    hl.append(r)
    paragraph._element.append(hl)


def add_styled_paragraph(doc, text, bold=False, italic=False, size=None, alignment=None, space_before=None, space_after=None):
    """Add a paragraph with optional styling."""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    if space_before is not None:
        p.paragraph_format.space_before = space_before
    if space_after is not None:
        p.paragraph_format.space_after = space_after

    # Parse inline formatting: **bold**, *italic*, `code`, [text](url), raw URLs
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\[[^\]]*\]\([^)\s]+\)|https?://[^\s]+)', text)
    for part in parts:
        if not part:
            continue
        link_match = re.match(r'^\[(.*?)\]\((https?://[^)\s]+)\)$', part)
        if link_match:
            add_hyperlink_run(p, link_match.group(2), link_match.group(1), size or FONT_SIZE, bold)
        elif part.startswith('http'):
            add_hyperlink_run(p, part.rstrip('.,;:'), part.rstrip('.,;:'), size or FONT_SIZE, bold)
        elif part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.name = FONT_NAME
            run.font.size = size or FONT_SIZE
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = p.add_run(part[1:-1])
            run.italic = True
            run.font.name = FONT_NAME
            run.font.size = size or FONT_SIZE
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        elif part.strip():
            run = p.add_run(part)
            run.font.name = FONT_NAME
            run.font.size = size or FONT_SIZE
            if bold:
                run.bold = True
            if italic:
                run.italic = True

    return p


def add_table_from_md(doc, lines, start_idx):
    """Parse a markdown table and add it to the document."""
    rows = []
    i = start_idx
    while i < len(lines) and lines[i].strip().startswith('|'):
        rows.append(lines[i].strip())
        i += 1

    if len(rows) < 2:
        return i

    # Parse cells from each row
    data = []
    for row in rows:
        cells = [c.strip() for c in row.strip('|').split('|')]
        data.append(cells)

    # Skip separator row (second row with :--- etc)
    header = data[0]
    body = [r for j, r in enumerate(data) if j != 1 and j != 0]

    num_cols = len(header)
    table = doc.add_table(rows=1 + len(body), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    # Header row
    for col_idx, cell_text in enumerate(header):
        cell_text = cell_text.replace('`', '')
        cell = table.rows[0].cells[col_idx]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(cell_text)
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, "1a3a5c")
        # White text for header
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Body rows
    for row_idx, row_data in enumerate(body):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < num_cols:
                cell_text = cell_text.replace('`', '')
                cell = table.rows[row_idx + 1].cells[col_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                # Parse bold markers in table cells
                if cell_text.startswith('**') and cell_text.endswith('**'):
                    run = p.add_run(cell_text[2:-2])
                    run.bold = True
                else:
                    run = p.add_run(cell_text)
                run.font.name = FONT_NAME
                run.font.size = Pt(10)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add spacing after table
    doc.add_paragraph()

    return i


def add_code_block(doc, code_text):
    """Add a styled code block to the document."""
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(50, 50, 50)

    doc.add_paragraph()  # spacing after code block


def parse_and_build(doc, toc_index):
    """Parse markdown content and build the document."""
    with open(MARKDOWN_PATH, 'r', encoding='utf-8') as f:
        content = f.read()

    # Skip metadata section (everything before first ---)
    parts = content.split('\n---\n', 1)
    body = parts[1] if len(parts) > 1 else parts[0]

    lines = body.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    code_lang = ""
    bookmark_counter = [0]  # separate counter for bookmark ids (must be unique)
    in_list = False

    while i < len(lines):
        line = lines[i]

        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                if code_lang == 'mermaid':
                    # Embed rendered mermaid PNG
                    if os.path.exists(ARCH_IMG_PATH):
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        run.add_picture(ARCH_IMG_PATH, width=Inches(5.5))
                        p2 = doc.add_paragraph()
                        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run2 = p2.add_run("Sekil 1: Sistem Mimari Katmanlari")
                        run2.font.name = FONT_NAME
                        run2.font.size = Pt(9)
                        run2.italic = True
                        run2.font.color.rgb = RGBColor(100, 100, 100)
                        doc.add_paragraph()
                    else:
                        add_code_block(doc, '\n'.join(code_lines))
                else:
                    add_code_block(doc, '\n'.join(code_lines))
                code_lines = []
                code_lang = ""
                in_code_block = False
            else:
                code_lang = line.strip()[3:].strip()
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Skip empty lines (handled as paragraph breaks)
        if not line.strip():
            if in_list:
                in_list = False
            i += 1
            continue

        # Horizontal rule → page break
        if line.strip() == '---':
            doc.add_page_break()
            i += 1
            continue

        # Heading 1 (#)
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            doc.add_heading(text, level=1)
            i += 1
            continue

        # Heading 2 (##)
        if line.startswith('## ') and not line.startswith('### '):
            text = line[3:].strip()
            heading_para = doc.add_heading(text, level=2)
            bid = bookmark_counter[0]
            add_heading_bookmark(heading_para, f"_Toc{toc_index[0]}", bid)
            toc_index[0] += 1
            bookmark_counter[0] = bid + 1
            i += 1
            continue

        # Heading 3 (###)
        if line.startswith('### '):
            text = line[4:].strip()
            heading_para = doc.add_heading(text, level=3)
            bid = bookmark_counter[0]
            add_heading_bookmark(heading_para, f"_Toc{toc_index[0]}", bid)
            toc_index[0] += 1
            bookmark_counter[0] = bid + 1
            i += 1
            continue

        # Tables
        if line.strip().startswith('|') and i + 1 < len(lines) and lines[i + 1].strip().startswith('|') and '---' in lines[i + 1]:
            i = add_table_from_md(doc, lines, i)
            continue

        # Bullet lists
        if re.match(r'^\s*[\-\*]\s+', line):
            text = re.sub(r'^\s*[\-\*]\s+', '', line)
            p = doc.add_paragraph(style='List Bullet')
            # Clear default text and add styled
            p.clear()
            _add_inline_runs(p, text, FONT_SIZE)
            i += 1
            continue

        # Numbered lists (1. 2. etc)
        if re.match(r'^\s*\d+\.\s+', line):
            text = re.sub(r'^\s*\d+\.\s+', '', line)
            p = doc.add_paragraph(style='List Number')
            p.clear()
            _add_inline_runs(p, text, FONT_SIZE)
            i += 1
            continue

        # Normal paragraph
        add_styled_paragraph(doc, line)
        i += 1


def _add_inline_runs(paragraph, text, font_size):
    """Parse inline formatting and add runs to a paragraph."""
    # Handle **bold** and *italic* and `code` and links
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\[[^\]]*\]\([^)\s]+\)|https?://[^\s]+)', text)
    for part in parts:
        if not part:
            continue
        link_match = re.match(r'^\[(.*?)\]\((https?://[^)\s]+)\)$', part)
        if link_match:
            add_hyperlink_run(paragraph, link_match.group(2), link_match.group(1), font_size)
        elif part.startswith('http'):
            add_hyperlink_run(paragraph, part.rstrip('.,;:'), part.rstrip('.,;:'), font_size)
        elif part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = FONT_NAME
            run.font.size = font_size
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.name = FONT_NAME
            run.font.size = font_size
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        elif part.strip():
            run = paragraph.add_run(part)
            run.font.name = FONT_NAME
            run.font.size = font_size


def main():
    global MARKDOWN_PATH, OUTPUT_PATH, COVER_TITLE, COVER_SUBTITLE, COVER_DATE

    parser = argparse.ArgumentParser(
        description="Markdown → DOCX converter with cover page, header, and embedded images."
    )
    parser.add_argument("--input", help="Input markdown file path")
    parser.add_argument("--output", help="Output .docx file path")
    parser.add_argument("--title", help="Cover page title (use \\n for line break)")
    parser.add_argument("--subtitle", help="Cover page subtitle")
    parser.add_argument("--date-line", help="Cover page date/version line")
    parser.add_argument("--compact", action="store_true", help="Tighter spacing (one-pager mode)")
    args = parser.parse_args()

    if args.input:
        MARKDOWN_PATH = os.path.abspath(args.input)
    if args.output:
        OUTPUT_PATH = os.path.abspath(args.output)
    if args.title:
        COVER_TITLE = args.title.replace("\\n", "\n")
    if args.subtitle is not None:
        COVER_SUBTITLE = args.subtitle
    if args.date_line is not None:
        COVER_DATE = args.date_line

    doc = Document()

    # ---- Global styles ----
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE
    style.paragraph_format.space_after = Pt(4 if args.compact else 6)
    style.paragraph_format.line_spacing = 1.05 if args.compact else 1.15

    # Set heading styles
    for i in range(1, 4):
        hs = doc.styles[f'Heading {i}']
        hs.font.name = FONT_NAME
        hs.font.color.rgb = RGBColor(0, 0, 0)
        sizes = {1: Pt(18), 2: Pt(15), 3: Pt(13)}
        hs.font.size = sizes.get(i, Pt(12))
        if args.compact:
            hs.paragraph_format.space_before = Pt(10 if i == 1 else 6)
            hs.paragraph_format.space_after = Pt(4)
        else:
            hs.paragraph_format.space_before = Pt(18 if i == 1 else 12)
            hs.paragraph_format.space_after = Pt(8)

    # Set list styles
    for list_style_name in ['List Bullet', 'List Number']:
        try:
            ls = doc.styles[list_style_name]
            ls.font.name = FONT_NAME
            ls.font.size = FONT_SIZE
            if args.compact:
                ls.paragraph_format.space_after = Pt(2)
        except KeyError:
            pass

    # ---- Page margins for all sections ----
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ---- COVER PAGE ----
    # First section = cover, no header
    first_section = doc.sections[0]
    first_section.different_first_page_header_footer = True
    # We add cover content immediately, then page break, then new section for content

    add_cover_page(doc)

    # Page break: cover → TOC
    doc.add_page_break()

    # ---- TABLE OF CONTENTS PAGE ----
    headings = extract_headings()
    add_toc_page(doc, headings)

    # ---- CONTENT SECTION with header ----
    # Add a section break after cover + TOC
    new_section = doc.add_section()
    new_section.different_first_page_header_footer = False
    new_section.top_margin = Cm(2)
    new_section.bottom_margin = Cm(2)
    new_section.left_margin = Cm(2.5)
    new_section.right_margin = Cm(2.5)

    # Add header to the content section
    add_header_image(new_section)

    # ---- Build document content ----
    toc_index = [0]  # shared counter for TOC anchor → content bookmark mapping
    parse_and_build(doc, toc_index)

    # ---- Save ----
    doc.save(OUTPUT_PATH)
    print(f"DOCX created: {OUTPUT_PATH}")


if __name__ == '__main__':
    main()
